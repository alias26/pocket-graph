"""Pass 3 extractors: PDF and Markdown -> graph nodes/edges.

Wraps the existing preprocessor_pdf_v2.py and preprocessor_md.py to emit
graph-compatible nodes/edges, so docs/papers can be ingested into the same
graph as code.
"""
from __future__ import annotations
import re
from pathlib import Path

from .extract import _make_id


def _slug(s: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "_", str(s)).strip("_").lower()[:60]


# ============================================================
# Markdown extractor
# ============================================================
def extract_markdown(path: Path, *, is_paper: bool = False) -> dict:
    """Extract a markdown/web-clip file into graph nodes/edges.

    If is_paper=True (file classified as PAPER by detect.py), only the
    file-root node is emitted. Heading-based section nodes ("Abstract",
    "Introduction", etc.) are skipped because they have no semantic value
    on their own — the LLM-driven `apply-semantic` step extracts the real
    concepts and adds them with proper section references.
    """
    text = path.read_text(encoding="utf-8", errors="replace")
    str_path = str(path)
    stem = path.stem
    file_nid = _make_id(stem)

    nodes: list[dict] = []
    edges: list[dict] = []

    # Frontmatter title (if present)
    title = path.stem
    fm_match = re.match(r"^---\s*\n(.*?)\n---\s*\n", text, re.S)
    if fm_match:
        t = re.search(r"^title:\s*(.+)$", fm_match.group(1), re.M)
        if t:
            title = t.group(1).strip().strip('"')

    # File node — papers are typed as 'paper', not 'document'
    nodes.append({
        "id": file_nid,
        "label": title,
        "file_type": "paper" if is_paper else "document",
        "source_file": str_path,
        "source_location": "L1",
    })

    # For papers, skip heading-based section nodes. The LLM apply-semantic
    # step will add the real concepts (transformer, multi-head-attention,
    # etc.) with proper section references. Heading-only nodes like
    # "Abstract" or "Introduction" carry no semantic value on their own
    # and clutter the god-node analysis.
    if is_paper:
        return {"nodes": nodes, "edges": edges}

    # Headings -> contained nodes
    used_anchors: dict[str, int] = {}
    headings: list[tuple[int, str, str, int]] = []  # (level, title, anchor, line)

    for m in re.finditer(r"^(#{1,6})\s+(.+?)\s*$", text, re.M):
        level = len(m.group(1))
        ttl = m.group(2).strip()
        anchor_base = _slug(ttl)
        if anchor_base in used_anchors:
            used_anchors[anchor_base] += 1
            anchor = f"{anchor_base}_{used_anchors[anchor_base]}"
        else:
            used_anchors[anchor_base] = 0
            anchor = anchor_base
        line = text[:m.start()].count("\n") + 1
        headings.append((level, ttl, anchor, line))

    parent_stack: list[tuple[int, str]] = []  # (level, nid)
    for level, ttl, anchor, line in headings:
        nid = _make_id(stem, anchor)
        nodes.append({
            "id": nid,
            "label": ttl,
            "file_type": "document",
            "source_file": str_path,
            "source_location": f"L{line}",
        })
        # Pop deeper or same-level ancestors
        while parent_stack and parent_stack[-1][0] >= level:
            parent_stack.pop()

        parent_nid = parent_stack[-1][1] if parent_stack else file_nid
        edges.append({
            "source": parent_nid,
            "target": nid,
            "relation": "contains",
            "confidence": "EXTRACTED",
            "source_file": str_path,
            "source_location": f"L{line}",
        })
        parent_stack.append((level, nid))

    # Code blocks
    for i, m in enumerate(re.finditer(r"^```(\S+)?\n(.*?)^```", text, re.M | re.S)):
        line = text[:m.start()].count("\n") + 1
        lang = m.group(1) or "text"
        code_snippet = m.group(2).strip()[:200]
        cb_nid = _make_id(stem, "codeblock", str(i + 1))
        nodes.append({
            "id": cb_nid,
            "label": f"[{lang}] {code_snippet[:50]}",
            "file_type": "document",
            "source_file": str_path,
            "source_location": f"L{line}",
            "language": lang,
            "body": m.group(2)[:1000],
        })
        edges.append({
            "source": file_nid,
            "target": cb_nid,
            "relation": "includes",
            "confidence": "EXTRACTED",
            "source_file": str_path,
            "source_location": f"L{line}",
        })

    # External links
    for m in re.finditer(r"(?<!!)\[([^\]]+)\]\((https?://[^)]+)\)", text):
        line = text[:m.start()].count("\n") + 1
        link_text = m.group(1)[:60]
        url = m.group(2)
        link_nid = _make_id(stem, "link", _slug(url)[:30])
        # Avoid duplicate link nodes
        if any(n["id"] == link_nid for n in nodes):
            continue
        nodes.append({
            "id": link_nid,
            "label": link_text,
            "file_type": "concept",
            "source_file": "",
            "source_location": "",
            "url": url,
        })
        edges.append({
            "source": file_nid,
            "target": link_nid,
            "relation": "references",
            "confidence": "EXTRACTED",
            "source_file": str_path,
            "source_location": f"L{line}",
        })

    return {"nodes": nodes, "edges": edges}


# ============================================================
# PDF extractor
# ============================================================
def extract_pdf(path: Path) -> dict:
    """Extract a PDF using preprocessor_pdf_v2 facts, mapped to graph nodes/edges."""
    # Use the v2.1 preprocessor (lives in outputs)
    import importlib.util
    pp_path = Path("/mnt/user-data/outputs/ingest_experiment/preprocessor_pdf_v2.py")
    if not pp_path.exists():
        pp_path = Path("/home/claude/preprocessor_pdf_v2.py")
    if not pp_path.exists():
        raise FileNotFoundError(
            "preprocessor_pdf_v2.py not found. Expected at outputs/ingest_experiment/."
        )
    spec = importlib.util.spec_from_file_location("preprocessor_pdf_v2", pp_path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)

    facts = mod.pdf_facts_v2(path)
    str_path = str(path)
    stem = path.stem
    file_nid = _make_id(stem)

    title = facts["metadata"].get("title", path.stem)
    nodes: list[dict] = []
    edges: list[dict] = []

    nodes.append({
        "id": file_nid,
        "label": title,
        "file_type": "paper",
        "source_file": str_path,
        "source_location": "p1",
        "page_count": facts["metadata"].get("page_count"),
        "author": facts["metadata"].get("author"),
    })

    # Sections (with section-id parent inference for X.Y.Z hierarchy)
    section_ids_by_path: dict[str, str] = {}  # "3.2" -> nid
    for s in facts["sections"]:
        title_clean = s["title"].strip()
        # Try to find a numeric prefix like "3.2.1"
        prefix_m = re.match(r"^(\d+(?:\.\d+){0,2})\b", title_clean)
        prefix = prefix_m.group(1) if prefix_m else _slug(title_clean)[:30]
        nid = _make_id(stem, "sec", prefix or _slug(title_clean))
        if nid in {n["id"] for n in nodes}:
            continue
        nodes.append({
            "id": nid,
            "label": title_clean,
            "file_type": "paper",
            "source_file": str_path,
            "source_location": f"p{s['page']}",
            "level": s.get("level"),
        })
        section_ids_by_path[prefix] = nid
        # Find parent: trim last .X
        if prefix and "." in prefix:
            parent_path = prefix.rsplit(".", 1)[0]
            if parent_path in section_ids_by_path:
                edges.append({
                    "source": section_ids_by_path[parent_path],
                    "target": nid,
                    "relation": "contains",
                    "confidence": "EXTRACTED",
                    "source_file": str_path,
                    "source_location": f"p{s['page']}",
                })
                continue
        # Otherwise file --contains--> top-level section
        edges.append({
            "source": file_nid,
            "target": nid,
            "relation": "contains",
            "confidence": "EXTRACTED",
            "source_file": str_path,
            "source_location": f"p{s['page']}",
        })

    # Figures
    for fig in facts.get("figures", []):
        if fig.get("source") == "image_object":
            continue  # skip image placeholders without captions
        fid = _make_id(stem, fig["id"])
        nodes.append({
            "id": fid,
            "label": f"{fig['id']}: {fig.get('caption','')[:80]}",
            "file_type": "paper",
            "source_file": str_path,
            "source_location": f"p{fig['page']}",
            "caption": fig.get("caption", ""),
        })
        edges.append({
            "source": file_nid,
            "target": fid,
            "relation": "includes",
            "confidence": "EXTRACTED",
            "source_file": str_path,
            "source_location": f"p{fig['page']}",
        })

    # Tables
    for tab in facts.get("tables", []):
        tid = _make_id(stem, tab["id"])
        nodes.append({
            "id": tid,
            "label": f"{tab['id']}: {tab.get('caption','')[:80]}",
            "file_type": "paper",
            "source_file": str_path,
            "source_location": f"p{tab['page']}",
            "caption": tab.get("caption", ""),
        })
        edges.append({
            "source": file_nid,
            "target": tid,
            "relation": "includes",
            "confidence": "EXTRACTED",
            "source_file": str_path,
            "source_location": f"p{tab['page']}",
        })

    return {"nodes": nodes, "edges": edges}


__all__ = ["extract_markdown", "extract_pdf", "extract_docx", "extract_xlsx"]


# ============================================================
# DOCX extractor
# ============================================================
def extract_docx(path: Path) -> dict:
    """Extract a Word document into graph nodes/edges.
    
    Captures: file, headings (Heading 1/2/3 styles), tables, hyperlinks.
    Body paragraphs are NOT chunked into nodes -- they're context for headings.
    """
    try:
        from docx import Document
    except ImportError:
        return {"nodes": [], "edges": [], "error": "python-docx not installed"}

    doc = Document(path)
    str_path = str(path)
    stem = path.stem
    file_nid = _make_id(stem)

    nodes: list[dict] = [{
        "id": file_nid,
        "label": path.name,
        "file_type": "document",
        "source_file": str_path,
        "source_location": "p1",
    }]
    edges: list[dict] = []

    parent_stack: list[tuple[int, str]] = []
    used_anchors: dict[str, int] = {}
    para_idx = 0

    for para in doc.paragraphs:
        para_idx += 1
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        # Headings only
        if style.startswith("Heading"):
            try:
                level = int(style.split(" ")[-1])
            except ValueError:
                level = 1
            anchor_base = _slug(text)[:40]
            if anchor_base in used_anchors:
                used_anchors[anchor_base] += 1
                anchor = f"{anchor_base}_{used_anchors[anchor_base]}"
            else:
                used_anchors[anchor_base] = 0
                anchor = anchor_base
            nid = _make_id(stem, anchor)
            nodes.append({
                "id": nid,
                "label": text[:120],
                "file_type": "document",
                "source_file": str_path,
                "source_location": f"para{para_idx}",
                "level": level,
            })
            while parent_stack and parent_stack[-1][0] >= level:
                parent_stack.pop()
            parent = parent_stack[-1][1] if parent_stack else file_nid
            edges.append({
                "source": parent, "target": nid,
                "relation": "contains", "confidence": "EXTRACTED",
                "source_file": str_path, "source_location": f"para{para_idx}",
            })
            parent_stack.append((level, nid))

    # Tables
    for i, table in enumerate(doc.tables):
        nid = _make_id(stem, "table", str(i + 1))
        first_row = " | ".join(cell.text.strip()[:30] for cell in table.rows[0].cells) if table.rows else ""
        nodes.append({
            "id": nid,
            "label": f"Table {i+1}: {first_row[:80]}",
            "file_type": "document",
            "source_file": str_path,
            "source_location": f"table{i+1}",
            "row_count": len(table.rows),
            "col_count": len(table.columns),
        })
        edges.append({
            "source": file_nid, "target": nid,
            "relation": "includes", "confidence": "EXTRACTED",
            "source_file": str_path, "source_location": f"table{i+1}",
        })

    return {"nodes": nodes, "edges": edges}


# ============================================================
# XLSX extractor
# ============================================================
def extract_xlsx(path: Path) -> dict:
    """Extract an Excel workbook into graph nodes/edges.
    
    Captures: workbook, sheets, named ranges, formulas referencing other sheets.
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"nodes": [], "edges": [], "error": "openpyxl not installed"}

    wb = load_workbook(path, data_only=False)
    str_path = str(path)
    stem = path.stem
    file_nid = _make_id(stem)

    nodes: list[dict] = [{
        "id": file_nid,
        "label": path.name,
        "file_type": "document",
        "source_file": str_path,
        "source_location": "wb",
    }]
    edges: list[dict] = []

    sheet_nids = {}
    for sheet_name in wb.sheetnames:
        sid = _make_id(stem, "sheet", _slug(sheet_name))
        sheet_nids[sheet_name] = sid
        ws = wb[sheet_name]
        nodes.append({
            "id": sid,
            "label": sheet_name,
            "file_type": "document",
            "source_file": str_path,
            "source_location": f"sheet:{sheet_name}",
            "max_row": ws.max_row,
            "max_col": ws.max_column,
        })
        edges.append({
            "source": file_nid, "target": sid,
            "relation": "contains", "confidence": "EXTRACTED",
            "source_file": str_path, "source_location": f"sheet:{sheet_name}",
        })

    # Cross-sheet references via formulas
    import re as _re
    SHEET_REF = _re.compile(r"(?:'([^']+)'|([A-Za-z_][A-Za-z0-9_ ]*))!")
    seen_refs: set[tuple[str, str]] = set()  # dedup (src_sheet, ref_sheet) pairs
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                v = cell.value
                if isinstance(v, str) and v.startswith("="):
                    refs = set()
                    for m in SHEET_REF.finditer(v):
                        ref = m.group(1) or m.group(2)
                        if ref:
                            refs.add(ref.strip())
                    for ref_sheet in refs:
                        if ref_sheet in sheet_nids and ref_sheet != sheet_name:
                            key = (sheet_name, ref_sheet)
                            if key in seen_refs:
                                continue
                            seen_refs.add(key)
                            edges.append({
                                "source": sheet_nids[sheet_name],
                                "target": sheet_nids[ref_sheet],
                                "relation": "references",
                                "confidence": "EXTRACTED",
                                "source_file": str_path,
                                "source_location": f"{sheet_name}!{cell.coordinate}",
                            })

    # Named ranges
    if hasattr(wb, "defined_names"):
        try:
            for name in wb.defined_names:
                nid = _make_id(stem, "namedrange", _slug(name))
                nodes.append({
                    "id": nid,
                    "label": f"NamedRange: {name}",
                    "file_type": "concept",
                    "source_file": "",
                    "source_location": "",
                })
                edges.append({
                    "source": file_nid, "target": nid,
                    "relation": "defines", "confidence": "EXTRACTED",
                    "source_file": str_path, "source_location": f"name:{name}",
                })
        except Exception:
            pass

    return {"nodes": nodes, "edges": edges}
